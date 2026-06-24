from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document(r'c:\krishna\python_learning\python files\Python_answer_doc.docx')

# Find all question paragraph indices (bold lines starting with Q)
question_indices = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and text.startswith('Q') and '. ' in text[:6]:
        if para.runs and para.runs[0].bold:
            question_indices.append(i)

print(f'Total questions: {len(question_indices)}')

def apply_shading(para):
    """Apply black background shading to a paragraph"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    
    existing_shd = pPr.find(qn('w:shd'))
    if existing_shd is not None:
        pPr.remove(existing_shd)
    
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="000000" w:themeFill="text1"/>')
    pPr.append(shd)

# Process questions from Q23 onwards (index 22 in 0-based)
styled_count = 0
for q_num in range(22, len(question_indices)):
    q_idx = question_indices[q_num]
    end_idx = question_indices[q_num + 1] if q_num + 1 < len(question_indices) else len(doc.paragraphs)
    
    # First non-empty paragraph after question is explanation - skip it
    explanation_found = False
    
    for j in range(q_idx + 1, end_idx):
        text = doc.paragraphs[j].text.strip()
        para = doc.paragraphs[j]
        
        # Skip if already has shading
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:shd')) is not None:
            continue
        
        if not explanation_found and text:
            explanation_found = True
            continue
        
        if explanation_found:
            apply_shading(para)
            styled_count += 1

print(f'Applied shading to {styled_count} paragraphs')

output_path = r'c:\krishna\python_learning\python files\Python_answer_doc_styled.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
